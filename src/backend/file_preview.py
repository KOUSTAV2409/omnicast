#!/usr/bin/env python3
"""Build a JSON preview payload for Omnicast FilePreviewView.

Usage:
  python3 file_preview.py <path>
"""
from __future__ import annotations

import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

TEXT_EXT = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json", ".jsonc",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".env", ".xml", ".html",
    ".htm", ".css", ".scss", ".less", ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs",
    ".py", ".pyi", ".rb", ".go", ".rs", ".c", ".h", ".cpp", ".hpp", ".cc", ".java",
    ".kt", ".swift", ".m", ".mm", ".sh", ".bash", ".zsh", ".fish", ".ps1", ".bat",
    ".qml", ".lua", ".sql", ".r", ".jl", ".php", ".pl", ".pm", ".vim", ".diff",
    ".patch", ".gitignore", ".dockerfile", ".makefile", ".cmake", ".nix",
}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".ico"}
OFFICE_EXT = {
    ".doc", ".docx", ".odt", ".rtf",
    ".xls", ".xlsx", ".ods", ".csv",
    ".ppt", ".pptx", ".odp",
}
PDF_EXT = {".pdf"}
MARKDOWN_EXT = {".md", ".markdown", ".mdown"}
CODE_EXT = {
    ".py", ".pyi", ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs", ".go", ".rs",
    ".c", ".h", ".cpp", ".hpp", ".cc", ".java", ".kt", ".swift", ".rb", ".php",
    ".sh", ".bash", ".zsh", ".fish", ".qml", ".lua", ".sql", ".css", ".scss",
    ".json", ".jsonc", ".yaml", ".yml", ".toml", ".nix", ".vim", ".diff", ".patch",
}
MAX_TEXT_BYTES = 200_000
MAX_TEXT_CHARS = 40_000
MAX_DIR_ENTRIES = 80
MAX_HIGHLIGHT_CHARS = 24_000

# Omarchy-ish syntax palette (readable on dark card)
C_FG = "#f5f0ec"
C_MUTED = "#6d7c8d"
C_ACCENT = "#fa8526"
C_STR = "#8fbf7a"
C_NUM = "#c9a0dc"
C_TYPE = "#6cb6ff"
C_PUNCT = "#9aa7b5"
C_BG_CODE = "#0e131b"

KEYWORDS = {
    ".py": {
        "def", "class", "return", "import", "from", "as", "if", "elif", "else",
        "for", "while", "try", "except", "finally", "with", "yield", "async",
        "await", "True", "False", "None", "and", "or", "not", "in", "is", "pass",
        "raise", "break", "continue", "lambda", "global", "nonlocal",
    },
    ".js": {
        "function", "const", "let", "var", "return", "import", "export", "from",
        "class", "extends", "if", "else", "for", "while", "try", "catch",
        "finally", "async", "await", "new", "this", "true", "false", "null",
        "typeof", "instanceof", "switch", "case", "break", "continue", "default",
    },
    ".ts": set(),  # filled below
    ".qml": {
        "property", "readonly", "signal", "function", "import", "id", "true",
        "false", "null", "undefined", "var", "if", "else", "for", "while",
        "return", "Component", "onCompleted",
    },
    ".rs": {
        "fn", "let", "mut", "pub", "struct", "enum", "impl", "trait", "use",
        "mod", "return", "if", "else", "match", "for", "while", "loop", "self",
        "Self", "true", "false", "async", "await", "crate", "super", "where",
    },
    ".go": {
        "func", "package", "import", "return", "if", "else", "for", "range",
        "switch", "case", "defer", "go", "chan", "select", "type", "struct",
        "interface", "map", "var", "const", "true", "false", "nil",
    },
}
KEYWORDS[".ts"] = KEYWORDS[".js"] | {"type", "interface", "implements", "readonly", "enum"}
KEYWORDS[".tsx"] = KEYWORDS[".ts"]
KEYWORDS[".jsx"] = KEYWORDS[".js"]
KEYWORDS[".sh"] = {
    "if", "then", "else", "fi", "for", "while", "do", "done", "case", "esac",
    "function", "return", "export", "local", "echo", "exit",
}
KEYWORDS[".bash"] = KEYWORDS[".sh"]
KEYWORDS[".zsh"] = KEYWORDS[".sh"]
KEYWORDS[".lua"] = {
    "function", "local", "return", "if", "then", "else", "elseif", "end", "for",
    "while", "do", "repeat", "until", "in", "and", "or", "not", "true", "false",
    "nil", "require",
}
KEYWORDS[".c"] = {
    "int", "char", "void", "return", "if", "else", "for", "while", "switch",
    "case", "break", "continue", "struct", "typedef", "enum", "const", "static",
    "sizeof", "include",
}
KEYWORDS[".cpp"] = KEYWORDS[".c"] | {
    "class", "public", "private", "protected", "namespace", "template", "typename",
    "using", "new", "delete", "this", "true", "false", "nullptr", "auto",
}
KEYWORDS[".java"] = {
    "class", "public", "private", "protected", "static", "void", "return", "if",
    "else", "for", "while", "new", "this", "import", "package", "extends",
    "implements", "try", "catch", "finally", "true", "false", "null",
}
KEYWORDS[".css"] = {
    "color", "background", "display", "flex", "grid", "margin", "padding",
    "border", "width", "height", "font", "position", "absolute", "relative",
    "important",
}
KEYWORDS[".sql"] = {
    "SELECT", "FROM", "WHERE", "INSERT", "UPDATE", "DELETE", "JOIN", "LEFT",
    "RIGHT", "INNER", "ON", "AND", "OR", "NOT", "NULL", "AS", "CREATE", "TABLE",
    "INDEX", "VALUES", "INTO", "SET", "ORDER", "BY", "GROUP", "LIMIT",
}
KEYWORDS[".php"] = {
    "function", "return", "if", "else", "elseif", "foreach", "for", "while",
    "class", "public", "private", "protected", "echo", "new", "true", "false",
    "null", "array",
}



def human_size(n: int) -> str:
    units = ["B", "KB", "MB", "GB", "TB"]
    size = float(n)
    for u in units:
        if size < 1024 or u == units[-1]:
            if u == "B":
                return f"{int(size)} {u}"
            return f"{size:.1f} {u}"
        size /= 1024
    return f"{n} B"


def mime_of(path: Path) -> str:
    # Prefer file(1) when available for accuracy
    file_bin = shutil.which("file")
    if file_bin:
        try:
            out = subprocess.check_output(
                [file_bin, "-b", "--mime-type", str(path)],
                text=True,
                stderr=subprocess.DEVNULL,
                timeout=2,
            ).strip()
            if out:
                return out
        except Exception:
            pass
    guess, _ = mimetypes.guess_type(str(path))
    return guess or "application/octet-stream"


def looks_binary(sample: bytes) -> bool:
    if b"\x00" in sample:
        return True
    # High ratio of non-text bytes
    if not sample:
        return False
    textish = sum(1 for b in sample if 9 <= b <= 13 or 32 <= b <= 126)
    return (textish / len(sample)) < 0.75


def read_text_file(path: Path) -> str:
    data = path.read_bytes()[:MAX_TEXT_BYTES]
    if looks_binary(data[:4096]):
        return ""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except Exception:
            text = ""
    if not text:
        return ""
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n… truncated …"
    return text


def extract_docx_text(path: Path) -> str:
    try:
        import docx  # type: ignore

        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts).strip()
        if text:
            if len(text) > MAX_TEXT_CHARS:
                text = text[:MAX_TEXT_CHARS] + "\n\n… truncated …"
            return text
    except Exception:
        pass

    # Fallback: unzip document.xml and strip tags
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        texts = [t.text for t in root.findall(".//w:t", ns) if t.text]
        text = " ".join(texts)
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n\n… truncated …"
        return text
    except Exception:
        return ""


def extract_pdf_text(path: Path) -> str:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        return ""
    try:
        out = subprocess.check_output(
            [pdftotext, "-l", "8", "-q", str(path), "-"],
            text=True,
            stderr=subprocess.DEVNULL,
            timeout=4,
        )
        text = out.strip()
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS] + "\n\n… truncated …"
        return text
    except Exception:
        return ""


def pdf_page_image(path: Path) -> str:
    """Render first PDF page to cache; return file URI or empty."""
    pdftoppm = shutil.which("pdftoppm")
    if not pdftoppm:
        return ""
    cache_dir = Path(os.environ.get("XDG_CACHE_HOME", Path.home() / ".cache")) / "omnicast"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out_base = cache_dir / "pdf-preview"
    # Clean previous render
    for suffix in (".png", "-1.png"):
        p = Path(str(out_base) + suffix) if suffix.startswith("-") else out_base.with_suffix(suffix)
        try:
            if p.exists():
                p.unlink()
        except Exception:
            pass
    try:
        subprocess.check_call(
            [pdftoppm, "-png", "-f", "1", "-l", "1", "-singlefile", "-r", "120", str(path), str(out_base)],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=6,
        )
    except Exception:
        return ""
    png = out_base.with_suffix(".png")
    if png.exists():
        return png.resolve().as_uri()
    return ""


def html_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _span(color: str, text: str) -> str:
    return f'<span style="color:{color}">{html_escape(text)}</span>'


def _lang_style(ext: str) -> dict:
    """Comment / string rules per language family."""
    hash_comments = ext in {".py", ".pyi", ".sh", ".bash", ".zsh", ".fish", ".rb", ".yml", ".yaml", ".toml", ".r", ".pl", ".pm"}
    slash_comments = ext in {
        ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs", ".go", ".rs", ".c", ".h", ".cpp",
        ".hpp", ".cc", ".java", ".kt", ".swift", ".qml", ".css", ".scss", ".less", ".php", ".sql",
    }
    return {
        "hash": hash_comments,
        "slash": slash_comments,
        "block": slash_comments or ext in {".css", ".scss", ".qml", ".java", ".c", ".cpp", ".h"},
        "backtick": ext in {".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs", ".qml"},
    }


def highlight_code_html(text: str, ext: str) -> str:
    """Editor-like syntax colors via a small tokenizer (no pygments)."""
    if len(text) > MAX_HIGHLIGHT_CHARS:
        text = text[:MAX_HIGHLIGHT_CHARS] + "\n… truncated …"

    if ext in {".json", ".jsonc"}:
        body = _highlight_json(text)
    else:
        body = _tokenize_highlight(text, ext)

    return (
        f'<div style="font-family:monospace; white-space:pre-wrap; '
        f'color:{C_FG}; font-size:13px; line-height:1.5; '
        f'background-color:{C_BG_CODE};">'
        f"{body}</div>"
    )


def _highlight_json(text: str) -> str:
    out: list[str] = []
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch in " \t\n\r":
            out.append(html_escape(ch))
            i += 1
            continue
        if ch in "{}[],:":
            out.append(_span(C_PUNCT, ch))
            i += 1
            continue
        if ch == '"':
            j = i + 1
            while j < n:
                if text[j] == "\\" and j + 1 < n:
                    j += 2
                    continue
                if text[j] == '"':
                    j += 1
                    break
                j += 1
            token = text[i:j]
            # key if next non-space is :
            k = j
            while k < n and text[k] in " \t\n\r":
                k += 1
            if k < n and text[k] == ":":
                out.append(_span(C_TYPE, token))
            else:
                out.append(_span(C_STR, token))
            i = j
            continue
        if ch.isdigit() or (ch == "-" and i + 1 < n and text[i + 1].isdigit()):
            j = i + 1
            while j < n and (text[j].isdigit() or text[j] in ".eE+-"):
                j += 1
            out.append(_span(C_NUM, text[i:j]))
            i = j
            continue
        if text.startswith("true", i) or text.startswith("false", i) or text.startswith("null", i):
            for lit in ("true", "false", "null"):
                if text.startswith(lit, i) and (i + len(lit) >= n or not text[i + len(lit)].isalnum()):
                    out.append(_span(C_ACCENT, lit))
                    i += len(lit)
                    break
            else:
                out.append(html_escape(ch))
                i += 1
            continue
        out.append(html_escape(ch))
        i += 1
    return "".join(out)


def _tokenize_highlight(text: str, ext: str) -> str:
    style = _lang_style(ext)
    keywords = KEYWORDS.get(ext) or set()
    if ext in {".mjs", ".cjs", ".jsx"}:
        keywords = KEYWORDS.get(".js", set())
    if ext == ".tsx":
        keywords = KEYWORDS.get(".ts", set())
    if ext in {".hpp", ".cc", ".hh"}:
        keywords = KEYWORDS.get(".cpp", set()) | KEYWORDS.get(".c", set())

    out: list[str] = []
    i = 0
    n = len(text)

    while i < n:
        ch = text[i]

        # whitespace
        if ch in " \t\n\r":
            out.append(ch if ch == "\n" else html_escape(ch))
            i += 1
            continue

        # hash comments
        if style["hash"] and ch == "#":
            j = i
            while j < n and text[j] != "\n":
                j += 1
            out.append(_span(C_MUTED, text[i:j]))
            i = j
            continue

        # line comments //
        if style["slash"] and text.startswith("//", i):
            j = i
            while j < n and text[j] != "\n":
                j += 1
            out.append(_span(C_MUTED, text[i:j]))
            i = j
            continue

        # block comments /* */
        if style["block"] and text.startswith("/*", i):
            j = text.find("*/", i + 2)
            j = n if j < 0 else j + 2
            out.append(_span(C_MUTED, text[i:j]))
            i = j
            continue

        # strings
        if ch in "'\"" or (style["backtick"] and ch == "`"):
            quote = ch
            j = i + 1
            while j < n:
                if text[j] == "\\" and j + 1 < n:
                    j += 2
                    continue
                if text[j] == quote:
                    j += 1
                    break
                # template ${ } — keep simple, stay in string
                j += 1
            out.append(_span(C_STR, text[i:j]))
            i = j
            continue

        # numbers
        if ch.isdigit() or (ch == "." and i + 1 < n and text[i + 1].isdigit()):
            j = i + 1
            while j < n and (text[j].isalnum() or text[j] in ".xXbBoO_"):
                j += 1
            out.append(_span(C_NUM, text[i:j]))
            i = j
            continue

        # identifiers / keywords
        if ch.isalpha() or ch == "_" or ch == "$":
            j = i + 1
            while j < n and (text[j].isalnum() or text[j] in "_$"):
                j += 1
            word = text[i:j]
            if word in keywords:
                out.append(_span(C_ACCENT, word))
            elif word[:1].isupper() and any(c.islower() for c in word[1:]):
                out.append(_span(C_TYPE, word))
            else:
                out.append(_span(C_FG, word))
            i = j
            continue

        # punctuation
        if ch in "{}[]().,;:?<>+-*/%=&|!~^":
            out.append(_span(C_PUNCT, ch))
            i += 1
            continue

        out.append(html_escape(ch))
        i += 1

    return "".join(out)


def render_markdown_html(text: str) -> str:
    """Zero-dep markdown → Qt-friendly RichText HTML."""
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n… truncated …"

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    parts: list[str] = [
        f'<div style="color:{C_FG}; font-size:14px; line-height:1.55;">'
    ]
    i = 0
    in_ul = False
    in_ol = False

    def close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            parts.append("</ul>")
            in_ul = False
        if in_ol:
            parts.append("</ol>")
            in_ol = False

    def inline_md(s: str) -> str:
        # order matters: code first, then bold/italic, links
        s = html_escape(s)

        def code_repl(m: re.Match) -> str:
            return (
                f'<span style="font-family:monospace; background-color:#1b2532; '
                f'color:{C_STR};">{m.group(1)}</span>'
            )

        s = re.sub(r"`([^`]+)`", code_repl, s)
        s = re.sub(
            r"\*\*([^*]+)\*\*",
            rf'<span style="font-weight:600; color:{C_FG};">\1</span>',
            s,
        )
        s = re.sub(
            r"(?<!\*)\*([^*]+)\*(?!\*)",
            rf'<span style="font-style:italic; color:{C_PUNCT};">\1</span>',
            s,
        )
        s = re.sub(
            r"\[([^\]]+)\]\(([^)]+)\)",
            rf'<a href="\2" style="color:{C_TYPE}; text-decoration:none;">\1</a>',
            s,
        )
        return s

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        fence = re.match(r"^```(\w+)?\s*$", stripped)
        if fence:
            close_lists()
            lang = (fence.group(1) or "").lower()
            ext_map = {
                "python": ".py", "py": ".py", "javascript": ".js", "js": ".js",
                "typescript": ".ts", "ts": ".ts", "tsx": ".tsx", "jsx": ".jsx",
                "rust": ".rs", "rs": ".rs", "go": ".go", "c": ".c", "cpp": ".cpp",
                "java": ".java", "lua": ".lua", "qml": ".qml", "bash": ".sh",
                "sh": ".sh", "shell": ".sh", "json": ".json", "yaml": ".yaml",
                "yml": ".yml", "toml": ".toml", "css": ".css", "sql": ".sql",
            }
            i += 1
            buf: list[str] = []
            while i < len(lines) and not re.match(r"^```\s*$", lines[i].strip()):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            code = "\n".join(buf)
            ext = ext_map.get(lang, ".txt")
            if lang:
                hi = _tokenize_highlight(code, ext) if ext != ".json" else _highlight_json(code)
            else:
                hi = html_escape(code)
            parts.append(
                f'<div style="font-family:monospace; white-space:pre-wrap; '
                f'background-color:{C_BG_CODE}; color:{C_FG}; font-size:12px; '
                f'line-height:1.45; margin:8px 0;">{hi}</div>'
            )
            continue

        if not stripped:
            close_lists()
            parts.append("<br/>")
            i += 1
            continue

        if stripped.startswith("---") and set(stripped) <= {"-", " "}:
            close_lists()
            parts.append(
                f'<hr style="border:none; border-top:1px solid {C_MUTED}; margin:10px 0;"/>'
            )
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            close_lists()
            level = len(heading.group(1))
            sizes = {1: 20, 2: 17, 3: 15, 4: 14, 5: 13, 6: 12}
            size = sizes.get(level, 14)
            color = C_ACCENT if level <= 2 else C_FG
            parts.append(
                f'<div style="font-size:{size}px; font-weight:600; color:{color}; '
                f'margin:10px 0 4px 0;">{inline_md(heading.group(2))}</div>'
            )
            i += 1
            continue

        if stripped.startswith("> "):
            close_lists()
            quote = stripped[2:]
            parts.append(
                f'<div style="border-left:3px solid {C_ACCENT}; padding-left:10px; '
                f'color:{C_MUTED}; margin:4px 0;">{inline_md(quote)}</div>'
            )
            i += 1
            continue

        ul = re.match(r"^[-*+]\s+(.*)$", stripped)
        if ul:
            if in_ol:
                parts.append("</ol>")
                in_ol = False
            if not in_ul:
                parts.append("<ul>")
                in_ul = True
            parts.append(
                f'<li style="color:{C_FG}; margin:2px 0;">{inline_md(ul.group(1))}</li>'
            )
            i += 1
            continue

        ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ol:
            if in_ul:
                parts.append("</ul>")
                in_ul = False
            if not in_ol:
                parts.append("<ol>")
                in_ol = True
            parts.append(
                f'<li style="color:{C_FG}; margin:2px 0;">{inline_md(ol.group(2))}</li>'
            )
            i += 1
            continue

        close_lists()
        parts.append(
            f'<div style="margin:3px 0;">{inline_md(stripped)}</div>'
        )
        i += 1

    close_lists()
    parts.append("</div>")
    return "".join(parts)


def list_dir(path: Path) -> list[dict]:
    entries = []
    try:
        children = sorted(path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
    except Exception:
        return []
    for child in children[:MAX_DIR_ENTRIES]:
        try:
            st = child.stat()
            entries.append(
                {
                    "name": child.name + ("/" if child.is_dir() else ""),
                    "path": str(child.resolve()),
                    "is_dir": child.is_dir(),
                    "size_label": "dir" if child.is_dir() else human_size(st.st_size),
                }
            )
        except Exception:
            continue
    return entries


def preferred_opener(path: Path, mime: str, ext: str) -> dict:
    """Suggest the best external viewer for this file."""
    onlyoffice = shutil.which("onlyoffice-desktopeditors")
    if ext in OFFICE_EXT and onlyoffice:
        return {
            "id": "onlyoffice",
            "title": "Open in OnlyOffice",
            "argv": [onlyoffice, "--view=" + str(path)],
        }
    if ext in PDF_EXT:
        return {"id": "xdg", "title": "Open with PDF viewer", "argv": ["xdg-open", str(path)]}
    if ext in IMAGE_EXT:
        return {"id": "xdg", "title": "Open with image viewer", "argv": ["xdg-open", str(path)]}
    return {"id": "xdg", "title": "Open with system app", "argv": ["xdg-open", str(path)]}


def preview(path_str: str) -> dict:
    raw = (path_str or "").strip()
    if not raw:
        return {"ok": False, "error": "No path provided"}
    path = Path(os.path.expanduser(raw)).resolve()
    if not path.exists():
        return {"ok": False, "error": f"Not found: {path}"}

    st = path.stat()
    base = {
        "ok": True,
        "path": str(path),
        "name": path.name,
        "subtitle": str(path).replace(str(Path.home()), "~", 1),
        "is_dir": path.is_dir(),
        "size": st.st_size,
        "size_label": "directory" if path.is_dir() else human_size(st.st_size),
        "modified": datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M"),
        "mime": "inode/directory" if path.is_dir() else mime_of(path),
        "ext": path.suffix.lower(),
        "text": "",
        "html": "",
        "text_format": "plain",
        "image": "",
        "entries": [],
        "kind": "binary",
        "opener": {"id": "xdg", "title": "Open with system app", "argv": ["xdg-open", str(path)]},
    }

    if path.is_dir():
        base["kind"] = "dir"
        base["entries"] = list_dir(path)
        base["opener"] = {
            "id": "xdg",
            "title": "Open folder",
            "argv": ["xdg-open", str(path)],
        }
        return base

    ext = path.suffix.lower()
    mime = base["mime"]
    base["opener"] = preferred_opener(path, mime, ext)

    if ext in IMAGE_EXT or mime.startswith("image/"):
        base["kind"] = "image"
        base["image"] = path.as_uri()
        return base

    if ext in PDF_EXT or mime == "application/pdf":
        base["kind"] = "pdf"
        img = pdf_page_image(path)
        if img:
            base["image"] = img
        base["text"] = extract_pdf_text(path) or (
            "" if img else "(No text layer extracted. Open externally to view.)"
        )
        return base

    if ext in {".docx"} or "wordprocessingml" in mime:
        base["kind"] = "docx"
        text = extract_docx_text(path)
        base["text"] = text or "(Could not extract text. Open in OnlyOffice to view.)"
        return base

    if ext in OFFICE_EXT or mime.startswith("application/vnd"):
        base["kind"] = "office"
        base["text"] = "Office document. Press Enter to open in OnlyOffice / system app."
        return base

    if ext in MARKDOWN_EXT:
        text = read_text_file(path)
        if text:
            base["kind"] = "markdown"
            base["text"] = text
            base["html"] = render_markdown_html(text)
            base["text_format"] = "html"
            return base

    if ext in CODE_EXT or ext in TEXT_EXT or mime.startswith("text/") or mime in (
        "application/json",
        "application/javascript",
        "application/xml",
        "application/x-sh",
        "application/toml",
    ):
        text = read_text_file(path)
        if text:
            if ext in CODE_EXT or ext in {".json", ".jsonc"}:
                base["kind"] = "code"
                base["text"] = text
                base["html"] = highlight_code_html(text, ext if ext else ".txt")
                base["text_format"] = "html"
            else:
                base["kind"] = "text"
                base["text"] = text
            return base

    # Sniff unknown small files as text
    if st.st_size <= MAX_TEXT_BYTES:
        text = read_text_file(path)
        if text:
            base["kind"] = "text"
            base["text"] = text
            return base

    base["kind"] = "binary"
    base["text"] = f"Binary file ({base['size_label']}, {mime}). Open externally to view."
    return base


def main() -> int:
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "Usage: file_preview.py <path> [--cache NAME]"}))
        return 2

    path_arg = sys.argv[1]
    cache_name = "file-preview.json"
    if len(sys.argv) >= 4 and sys.argv[2] == "--cache":
        cache_name = sys.argv[3]
    elif len(sys.argv) >= 3 and sys.argv[2].startswith("--cache="):
        cache_name = sys.argv[2].split("=", 1)[1]

    payload = preview(path_arg)
    cache_dir = Path(os.environ.get("XDG_CACHE_HOME", Path.home() / ".cache")) / "omnicast"
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / cache_name
    cache_file.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    print(
        json.dumps(
            {
                "ok": bool(payload.get("ok")),
                "kind": payload.get("kind", ""),
                "path": payload.get("path", path_arg),
                "cache": str(cache_file),
                "error": payload.get("error", ""),
            },
            ensure_ascii=False,
        )
    )
    return 0 if payload.get("ok") else 1


if __name__ == "__main__":
    sys.exit(main())
